#!/usr/bin/env python3
"""Generate DOCX and PDF publication of the Structured Products Desk Bible."""

import re
import os
import sys
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC = os.path.join(BASE, "release", "desk-bible-v1", "Desk_Bible_v1.md")
OUT_DIR = os.path.join(BASE, "release", "desk-bible-v1")

RELEASE_TAG = "desk-bible-v1-publication"
VERSION = "1.0"
DATE = "2026-06-19"


def sanitize_for_pdf(text):
    """Replace Unicode characters that core PDF fonts can't handle."""
    replacements = {
        '—': '--',   # em dash
        '–': '-',    # en dash
        '‘': "'",    # left single quote
        '’': "'",    # right single quote
        '“': '"',    # left double quote
        '”': '"',    # right double quote
        '…': '...',  # ellipsis
        '•': '*',    # bullet
        '→': '->',   # right arrow
        '←': '<-',   # left arrow
        '≈': '~',    # approximately
        '±': '+/-',  # plus-minus
        '≥': '>=',   # gte
        '≤': '<=',   # lte
        '≠': '!=',   # not equal
        '─': '-',    # box drawing
        '━': '-',    # heavy box drawing
        '┃': '|',    # heavy vertical
        '█': '#',    # full block
        '●': '*',    # black circle
        '✓': 'Y',    # check mark
        '✗': 'X',    # ballot X
        '☐': '[ ]',  # ballot box
        '☑': '[x]',  # checked box
        '×': 'x',    # multiplication sign
        ' ': ' ',    # thin space
        '​': '',     # zero-width space
        ' ': ' ',    # non-breaking space
        '−': '-',    # minus sign
        'α': 'alpha',
        'β': 'beta',
        'γ': 'gamma',
        'δ': 'delta',
        'σ': 'sigma',
        'ρ': 'rho',
        'μ': 'mu',
        'Δ': 'Delta',
        'Σ': 'Sigma',
        '∙': '*',    # bullet operator
        '―': '--',   # horizontal bar
        '‐': '-',    # hyphen
        '‑': '-',    # non-breaking hyphen
        '‒': '-',    # figure dash
        '‚': ',',    # single low-9 quote
        '„': '"',    # double low-9 quote
        '′': "'",    # prime
        '″': '"',    # double prime
        '▶': '>',    # right-pointing triangle
        '▸': '>',    # right-pointing small triangle
        '─': '-',    # box drawings light horizontal
        '═': '=',    # box drawings double horizontal
        '│': '|',    # box drawings light vertical
        '┌': '+',    # box drawings light down and right
        '┐': '+',    # box drawings light down and left
        '└': '+',    # box drawings light up and right
        '┘': '+',    # box drawings light up and left
        '┼': '+',    # box drawings light vertical and horizontal
        '┴': '+',    # box drawings light up and horizontal
        '┬': '+',    # box drawings light down and horizontal
        '├': '+',    # box drawings light vertical and right
        '┤': '+',    # box drawings light vertical and left
        '▁': '_',    # lower one eighth block
        '▄': '_',    # lower half block
        '▀': '^',    # upper half block
        '■': '#',    # black square
        '□': '[ ]',  # white square
        '○': 'o',    # white circle
        '◎': '(o)',  # bullseye
        '◆': '<>',   # black diamond
        '◇': '<>',   # white diamond
        '★': '*',    # black star
        '☆': '*',    # white star
        '↑': '^',    # up arrow
        '↓': 'v',    # down arrow
        '⇒': '=>',   # double right arrow
        '⇐': '<=',   # double left arrow
        '²': '2',    # superscript 2
        '³': '3',    # superscript 3
        '¹': '1',    # superscript 1
        '⁰': '0',    # superscript 0
        '⁴': '4',    # superscript 4
        '½': '1/2',  # fraction 1/2
        '¼': '1/4',  # fraction 1/4
        '¾': '3/4',  # fraction 3/4
        'é': 'e',    # e acute
        'è': 'e',    # e grave
        'à': 'a',    # a grave
        'ü': 'u',    # u umlaut
        'ö': 'o',    # o umlaut
        'ä': 'a',    # a umlaut
        'ñ': 'n',    # n tilde
    }
    for char, replacement in replacements.items():
        text = text.replace(char, replacement)
    # Catch any remaining non-latin1 characters
    result = []
    for ch in text:
        try:
            ch.encode('latin-1')
            result.append(ch)
        except UnicodeEncodeError:
            result.append('?')
    return ''.join(result)

DARK = RGBColor(0x1B, 0x2A, 0x4A)
ACCENT = RGBColor(0x2E, 0x75, 0xB6)
GREY = RGBColor(0x58, 0x58, 0x58)
LIGHT_BG = "D9E2F3"
TABLE_HEADER_BG = "1B2A4A"
TABLE_ALT_BG = "F2F2F2"


def parse_markdown(path):
    with open(path, "r") as f:
        lines = f.readlines()
    return lines


def set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    tcPr.append(shading)


def add_header_footer(doc, section_obj):
    header = section_obj.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("Structured Products Desk Bible  |  v1.0  |  CONFIDENTIAL")
    run.font.size = Pt(7.5)
    run.font.color.rgb = GREY
    run.font.italic = True

    footer = section_obj.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = GREY
    fld_char1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fld_char1)
    run2 = fp.add_run()
    run2.font.size = Pt(8)
    run2.font.color.rgb = GREY
    instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instr)
    run3 = fp.add_run()
    fld_char2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fld_char2)


def add_cover_page(doc):
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("STRUCTURED PRODUCTS")
    run.font.size = Pt(36)
    run.font.color.rgb = DARK
    run.bold = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("DESK BIBLE")
    run2.font.size = Pt(36)
    run2.font.color.rgb = DARK
    run2.bold = True

    doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("━" * 40)
    run3.font.size = Pt(14)
    run3.font.color.rgb = ACCENT

    doc.add_paragraph()

    details = [
        ("Version", VERSION),
        ("Date", DATE),
        ("Release Tag", RELEASE_TAG),
        ("Products", "28"),
        ("Families", "ELN (13), SRT (5), CLN (5), STEG (4), Swap (1)"),
        ("Classification", "INTERNAL — CONFIDENTIAL"),
    ]
    for label, value in details:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_l = p.add_run(f"{label}: ")
        run_l.font.size = Pt(12)
        run_l.font.color.rgb = GREY
        run_v = p.add_run(value)
        run_v.font.size = Pt(12)
        run_v.font.color.rgb = DARK
        run_v.bold = True

    for _ in range(4):
        doc.add_paragraph()

    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_foot = p_foot.add_run("Generated from: release/desk-bible-v1/Desk_Bible_v1.md")
    run_foot.font.size = Pt(9)
    run_foot.font.color.rgb = GREY
    run_foot.italic = True

    doc.add_page_break()


def add_toc(doc):
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p_title.add_run("TABLE OF CONTENTS")
    run.font.size = Pt(20)
    run.font.color.rgb = DARK
    run.bold = True

    doc.add_paragraph()

    parts = {
        "Part 3 — Equity-Linked Notes": [
            ("Section 3.1", "Reverse Convertibles", [
                "Reverse Convertible", "Discounted Reverse Convertible",
                "Knock-Out Discounted Reverse Convertible", "Callable Reverse Convertible",
                "Airbag / Leveraged Reverse Convertible"
            ]),
            ("Section 3.2", "Autocallables", [
                "Phoenix Autocallable", "Fixed Coupon Note (FCN)",
                "Callable Range Accrual ELN"
            ]),
            ("Section 3.3", "Other ELN Structures", [
                "Bonus / Participation Note", "Principal Protected Note",
                "Warrant / Turbo Certificate", "Issuer Callable Note",
                "Digital / Exotic Coupon Notes"
            ]),
        ],
        "Part 4 — Structured Rate Trades": [
            (None, None, [
                "IR Callable Fixed Rate Note", "Accreting Fixed Rate Note",
                "Non-Callable Range Accrual", "Callable Range Accrual SRT",
                "Digital Cap-Floor Note"
            ]),
        ],
        "Part 5 — Swaps": [
            (None, None, ["Vanilla Interest Rate Swap"]),
        ],
        "Part 6 — Credit-Linked Notes": [
            (None, None, [
                "Vanilla Credit-Linked Note", "Skew CLN",
                "First-to-Default (FTD)", "Nth-to-Default (NTD)",
                "Synthetic CDO Tranche (Tranche)"
            ]),
        ],
        "Part 7 — Steepener Notes": [
            (None, None, [
                "Vanilla Steepener Note", "Range-Accrual Steepener",
                "Callable Steepener", "Target Accrual Redemption Note (TARN)"
            ]),
        ],
    }

    for part_name, sections in parts.items():
        p = doc.add_paragraph()
        run = p.add_run(part_name)
        run.font.size = Pt(13)
        run.font.color.rgb = DARK
        run.bold = True
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(2)

        for sec_num, sec_name, products in sections:
            if sec_num:
                ps = doc.add_paragraph()
                ps.paragraph_format.left_indent = Cm(1)
                ps.paragraph_format.space_before = Pt(4)
                ps.paragraph_format.space_after = Pt(1)
                run_s = ps.add_run(f"{sec_num} — {sec_name}")
                run_s.font.size = Pt(11)
                run_s.font.color.rgb = ACCENT
                run_s.bold = True

            for product in products:
                pp = doc.add_paragraph()
                pp.paragraph_format.left_indent = Cm(2)
                pp.paragraph_format.space_before = Pt(1)
                pp.paragraph_format.space_after = Pt(1)
                run_p = pp.add_run(f"•  {product}")
                run_p.font.size = Pt(10)
                run_p.font.color.rgb = GREY

    doc.add_paragraph()

    appendices_title = doc.add_paragraph()
    run_a = appendices_title.add_run("Appendices")
    run_a.font.size = Pt(13)
    run_a.font.color.rgb = DARK
    run_a.bold = True

    for item in ["Release Notes Reference", "Product Catalog", "Quality Metrics Summary"]:
        pa = doc.add_paragraph()
        pa.paragraph_format.left_indent = Cm(1)
        run_ai = pa.add_run(f"•  {item}")
        run_ai.font.size = Pt(10)
        run_ai.font.color.rgb = GREY

    doc.add_page_break()


def process_inline(paragraph, text, default_bold=False, default_size=Pt(10)):
    """Handle **bold** and *italic* inline formatting."""
    pattern = re.compile(r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))')
    for m in pattern.finditer(text):
        if m.group(2):  # ***bold italic***
            run = paragraph.add_run(m.group(2))
            run.bold = True
            run.italic = True
        elif m.group(3):  # **bold**
            run = paragraph.add_run(m.group(3))
            run.bold = True
        elif m.group(4):  # *italic*
            run = paragraph.add_run(m.group(4))
            run.italic = True
        elif m.group(5):  # plain text
            run = paragraph.add_run(m.group(5))
            if default_bold:
                run.bold = True
        else:
            continue
        run.font.size = default_size
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_table(doc, table_lines):
    rows_raw = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows_raw.append(cells)

    if len(rows_raw) < 2:
        return

    header = rows_raw[0]
    separator_idx = None
    for i, row in enumerate(rows_raw):
        if all(re.match(r'^[-:]+$', c) for c in row):
            separator_idx = i
            break

    if separator_idx is not None:
        data_rows = rows_raw[separator_idx + 1:]
    else:
        data_rows = rows_raw[1:]

    num_cols = len(header)
    table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    for i, cell_text in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(cell_text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, TABLE_HEADER_BG)

    for r_idx, row_data in enumerate(data_rows):
        for c_idx in range(min(len(row_data), num_cols)):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            process_inline(p, row_data[c_idx], default_size=Pt(9))
            if r_idx % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_BG)

    doc.add_paragraph()


def add_appendices(doc):
    doc.add_page_break()

    p_title = doc.add_paragraph()
    run = p_title.add_run("APPENDICES")
    run.font.size = Pt(22)
    run.font.color.rgb = DARK
    run.bold = True

    doc.add_paragraph()

    # Appendix A
    pa = doc.add_paragraph()
    run_a = pa.add_run("Appendix A — Release Notes Reference")
    run_a.font.size = Pt(14)
    run_a.font.color.rgb = ACCENT
    run_a.bold = True

    doc.add_paragraph()
    info = [
        ("Release Tag", "desk-bible-v1-publication"),
        ("Source Document", "release/desk-bible-v1/Desk_Bible_v1.md"),
        ("Release Notes", "release/desk-bible-v1/release_notes_v1.md"),
        ("Final Report", "release/desk-bible-v1/Desk_Bible_Final_Report.md"),
        ("Maintenance Guide", "release/desk-bible-v1/Maintenance_and_Governance_Guide.md"),
        ("Project Closeout", "reports/Project_Closeout_Report.md"),
    ]
    for label, value in info:
        p = doc.add_paragraph()
        run_l = p.add_run(f"{label}: ")
        run_l.bold = True
        run_l.font.size = Pt(10)
        run_v = p.add_run(value)
        run_v.font.size = Pt(10)
        run_v.font.color.rgb = GREY

    doc.add_paragraph()

    # Appendix B
    pb = doc.add_paragraph()
    run_b = pb.add_run("Appendix B — Product Catalog")
    run_b.font.size = Pt(14)
    run_b.font.color.rgb = ACCENT
    run_b.bold = True

    doc.add_paragraph()

    products = [
        ("ELN", "RC001", "Reverse Convertible"),
        ("ELN", "DRC001", "Discounted Reverse Convertible"),
        ("ELN", "KODRC001", "Knock-Out Discounted Reverse Convertible"),
        ("ELN", "CRC001", "Callable Reverse Convertible"),
        ("ELN", "AIRBAG001", "Airbag / Leveraged Reverse Convertible"),
        ("ELN", "PHX001", "Phoenix Autocallable"),
        ("ELN", "FCN001", "Fixed Coupon Note"),
        ("ELN", "CRAELN001", "Callable Range Accrual ELN"),
        ("ELN", "BONUS001", "Bonus / Participation Note"),
        ("ELN", "PPN001", "Principal Protected Note"),
        ("ELN", "WARRANT001", "Warrant / Turbo Certificate"),
        ("ELN", "ICN001", "Issuer Callable Note"),
        ("ELN", "DIGITAL001", "Digital / Exotic Coupon Notes"),
        ("SRT", "IRCFRN001", "IR Callable Fixed Rate Note"),
        ("SRT", "AFRN001", "Accreting Fixed Rate Note"),
        ("SRT", "NCRA001", "Non-Callable Range Accrual"),
        ("SRT", "CRASRT001", "Callable Range Accrual SRT"),
        ("SRT", "DCFN001", "Digital Cap-Floor Note"),
        ("Swap", "SWAP001", "Vanilla Interest Rate Swap"),
        ("CLN", "VCLN001", "Vanilla Credit-Linked Note"),
        ("CLN", "SCLN001", "Skew CLN"),
        ("CLN", "FTD001", "First-to-Default"),
        ("CLN", "NTD001", "Nth-to-Default"),
        ("CLN", "TRANCHE001", "Synthetic CDO Tranche"),
        ("STEG", "VSTEG001", "Vanilla Steepener Note"),
        ("STEG", "RASTEG001", "Range-Accrual Steepener"),
        ("STEG", "CSTEG001", "Callable Steepener"),
        ("STEG", "TARN001", "Target Accrual Redemption Note"),
    ]

    table = doc.add_table(rows=1 + len(products), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    for i, h in enumerate(["#", "Family", "ID", "Product Name"]):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, TABLE_HEADER_BG)

    for idx, (family, pid, name) in enumerate(products):
        row = table.rows[idx + 1]
        for ci, val in enumerate([str(idx + 1), family, pid, name]):
            cell = row.cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9)
            if idx % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_BG)

    doc.add_paragraph()

    # Appendix C
    pc = doc.add_paragraph()
    run_c = pc.add_run("Appendix C — Quality Metrics Summary")
    run_c.font.size = Pt(14)
    run_c.font.color.rgb = ACCENT
    run_c.bold = True

    doc.add_paragraph()

    metrics = [
        ("Accuracy BLOCKERs", "0"),
        ("Style MUST_FIX", "0"),
        ("Broken Cross-References", "0"),
        ("Publishing Failures", "0"),
        ("Gate Pass Rate", "100% (all 3 gates, all 28 products)"),
        ("Retries (Batch 1+)", "0 across 22 products"),
        ("Memory Artifacts", "15 created, 69 reused (4.6:1 ratio)"),
        ("False Positives Suppressed", "84"),
        ("Tokens Per Product", "~14,866 (stable ±0%)"),
        ("Total Tokens", "~416,248"),
    ]

    table2 = doc.add_table(rows=1 + len(metrics), cols=2)
    table2.alignment = WD_TABLE_ALIGNMENT.LEFT
    table2.style = 'Table Grid'

    for i, h in enumerate(["Metric", "Value"]):
        cell = table2.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, TABLE_HEADER_BG)

    for idx, (metric, value) in enumerate(metrics):
        row = table2.rows[idx + 1]
        for ci, val in enumerate([metric, value]):
            cell = row.cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9)
            if ci == 1:
                run.bold = True
            if idx % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_BG)


def generate_docx(lines, out_path):
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    add_header_footer(doc, section)
    add_cover_page(doc)
    add_toc(doc)

    i = 0
    table_buffer = []
    in_table = False
    first_part = True
    skip_meta = True

    while i < len(lines):
        line = lines[i].rstrip("\n")

        if skip_meta:
            if line.startswith("# ") or line.startswith("**") or line == "---" or line == "":
                i += 1
                continue
            if line.startswith("## "):
                skip_meta = False
            else:
                i += 1
                continue

        if in_table:
            if line.startswith("|"):
                table_buffer.append(line)
                i += 1
                continue
            else:
                add_table(doc, table_buffer)
                table_buffer = []
                in_table = False

        if line.startswith("|") and "|" in line[1:]:
            in_table = True
            table_buffer = [line]
            i += 1
            continue

        if line == "---":
            i += 1
            continue

        if line.startswith("## "):
            heading_text = line[3:].strip()
            if not first_part:
                doc.add_page_break()
            first_part = False

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(heading_text.upper())
            run.font.size = Pt(22)
            run.font.color.rgb = DARK
            run.bold = True

            p_line = doc.add_paragraph()
            run_line = p_line.add_run("━" * 60)
            run_line.font.size = Pt(10)
            run_line.font.color.rgb = ACCENT

            i += 1
            continue

        if line.startswith("### Section "):
            heading_text = line[4:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(heading_text)
            run.font.size = Pt(14)
            run.font.color.rgb = ACCENT
            run.bold = True
            i += 1
            continue

        if line.startswith("### "):
            heading_text = line[4:].strip()
            doc.add_page_break()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(heading_text)
            run.font.size = Pt(16)
            run.font.color.rgb = DARK
            run.bold = True

            p_line = doc.add_paragraph()
            run_line = p_line.add_run("─" * 50)
            run_line.font.size = Pt(8)
            run_line.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

            i += 1
            continue

        if line.startswith("#### "):
            heading_text = line[5:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(heading_text)
            run.font.size = Pt(12)
            run.font.color.rgb = ACCENT
            run.bold = True
            i += 1
            continue

        if line.startswith("- ") or line.startswith("* "):
            bullet_text = line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run_bullet = p.add_run("•  ")
            run_bullet.font.size = Pt(10)
            run_bullet.font.color.rgb = ACCENT
            process_inline(p, bullet_text)
            i += 1
            continue

        if line.strip() == "":
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        process_inline(p, line)
        i += 1

    if in_table and table_buffer:
        add_table(doc, table_buffer)

    add_appendices(doc)

    doc.save(out_path)
    print(f"DOCX saved: {out_path}")
    print(f"DOCX size: {os.path.getsize(out_path):,} bytes")


def generate_pdf(lines, out_path):
    from fpdf import FPDF

    class DeskBiblePDF(FPDF):
        def __init__(self):
            super().__init__()
            self.set_auto_page_break(auto=True, margin=25)
            self._in_body = False

        def cell(self, w=0, h=None, text="", **kwargs):
            return super().cell(w, h, sanitize_for_pdf(str(text)), **kwargs)

        def multi_cell(self, w, h=None, text="", **kwargs):
            return super().multi_cell(w, h, sanitize_for_pdf(str(text)), **kwargs)

        def get_string_width(self, s):
            return super().get_string_width(sanitize_for_pdf(str(s)))

        def header(self):
            if not self._in_body:
                return
            self.set_font("Helvetica", "I", 7)
            self.set_text_color(88, 88, 88)
            super().cell(0, 5, sanitize_for_pdf("Structured Products Desk Bible  |  v1.0  |  CONFIDENTIAL"), align="R")
            self.ln(8)

        def footer(self):
            if not self._in_body:
                return
            self.set_y(-20)
            self.set_font("Helvetica", "", 8)
            self.set_text_color(88, 88, 88)
            super().cell(0, 10, f"Page {self.page_no()}", align="C")

    pdf = DeskBiblePDF()
    pdf.set_margins(25, 20, 25)

    # Cover page
    pdf.add_page()
    pdf.ln(60)
    pdf.set_font("Helvetica", "B", 36)
    pdf.set_text_color(27, 42, 74)
    pdf.cell(0, 16, "STRUCTURED PRODUCTS", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 16, "DESK BIBLE", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(10)
    pdf.set_draw_color(46, 117, 182)
    pdf.set_line_width(0.5)
    x_start = 50
    x_end = pdf.w - 50
    pdf.line(x_start, pdf.get_y(), x_end, pdf.get_y())
    pdf.ln(15)

    details = [
        ("Version", VERSION),
        ("Date", DATE),
        ("Release Tag", RELEASE_TAG),
        ("Products", "28"),
        ("Families", "ELN (13), SRT (5), CLN (5), STEG (4), Swap (1)"),
        ("Classification", "INTERNAL — CONFIDENTIAL"),
    ]
    for label, value in details:
        pdf.set_font("Helvetica", "", 11)
        pdf.set_text_color(88, 88, 88)
        label_w = pdf.get_string_width(f"{label}: ")
        pdf.set_font("Helvetica", "B", 11)
        value_w = pdf.get_string_width(value)
        total_w = label_w + value_w
        x = (pdf.w - total_w) / 2
        pdf.set_x(x)
        pdf.set_font("Helvetica", "", 11)
        pdf.set_text_color(88, 88, 88)
        pdf.cell(label_w, 8, f"{label}: ", new_x="END")
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(27, 42, 74)
        pdf.cell(value_w, 8, value, new_x="LMARGIN", new_y="NEXT")

    pdf.ln(30)
    pdf.set_font("Helvetica", "I", 9)
    pdf.set_text_color(88, 88, 88)
    pdf.cell(0, 6, "Generated from: release/desk-bible-v1/Desk_Bible_v1.md", align="C")

    # TOC page
    pdf._in_body = True
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 20)
    pdf.set_text_color(27, 42, 74)
    pdf.cell(0, 12, "TABLE OF CONTENTS", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(6)

    toc_parts = [
        ("Part 3 — Equity-Linked Notes", [
            "  Section 3.1 — Reverse Convertibles",
            "    Reverse Convertible", "    Discounted Reverse Convertible",
            "    Knock-Out Discounted RC", "    Callable Reverse Convertible",
            "    Airbag / Leveraged RC",
            "  Section 3.2 — Autocallables",
            "    Phoenix Autocallable", "    Fixed Coupon Note (FCN)",
            "    Callable Range Accrual ELN",
            "  Section 3.3 — Other ELN Structures",
            "    Bonus / Participation Note", "    Principal Protected Note",
            "    Warrant / Turbo Certificate", "    Issuer Callable Note",
            "    Digital / Exotic Coupon Notes",
        ]),
        ("Part 4 — Structured Rate Trades", [
            "    IR Callable Fixed Rate Note", "    Accreting Fixed Rate Note",
            "    Non-Callable Range Accrual", "    Callable Range Accrual SRT",
            "    Digital Cap-Floor Note",
        ]),
        ("Part 5 — Swaps", ["    Vanilla Interest Rate Swap"]),
        ("Part 6 — Credit-Linked Notes", [
            "    Vanilla CLN", "    Skew CLN",
            "    First-to-Default (FTD)", "    Nth-to-Default (NTD)",
            "    Synthetic CDO Tranche",
        ]),
        ("Part 7 — Steepener Notes", [
            "    Vanilla Steepener Note", "    Range-Accrual Steepener",
            "    Callable Steepener", "    TARN Steepener",
        ]),
    ]

    for part_title, items in toc_parts:
        pdf.set_font("Helvetica", "B", 12)
        pdf.set_text_color(27, 42, 74)
        pdf.cell(0, 8, part_title, new_x="LMARGIN", new_y="NEXT")
        for item in items:
            stripped = item.lstrip()
            indent = len(item) - len(stripped)
            if stripped.startswith("Section"):
                pdf.set_font("Helvetica", "B", 10)
                pdf.set_text_color(46, 117, 182)
                pdf.set_x(25 + indent * 2)
            else:
                pdf.set_font("Helvetica", "", 9)
                pdf.set_text_color(88, 88, 88)
                pdf.set_x(25 + indent * 2)
            pdf.cell(0, 5.5, stripped, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

    pdf.set_font("Helvetica", "B", 12)
    pdf.set_text_color(27, 42, 74)
    pdf.cell(0, 8, "Appendices", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(88, 88, 88)
    for a in ["Release Notes Reference", "Product Catalog", "Quality Metrics Summary"]:
        pdf.set_x(35)
        pdf.cell(0, 5.5, f"•  {a}", new_x="LMARGIN", new_y="NEXT")

    # Body content
    i = 0
    skip_meta = True
    in_table = False
    table_buffer = []
    effective_width = pdf.w - 50

    def write_paragraph(text, font_family="Helvetica", font_style="", font_size=10,
                        text_color=(51, 51, 51), indent=0):
        clean = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', text)
        clean = re.sub(r'\*\*(.+?)\*\*', r'\1', clean)
        clean = re.sub(r'\*(.+?)\*', r'\1', clean)
        clean = re.sub(r'---', '—', clean)

        pdf.set_font(font_family, font_style, font_size)
        pdf.set_text_color(*text_color)
        if indent:
            pdf.set_x(25 + indent)
        pdf.multi_cell(effective_width - indent, 5, clean)
        pdf.ln(1)

    def render_table(tbl_lines):
        rows = []
        for tl in tbl_lines:
            cells = [c.strip() for c in tl.strip().strip("|").split("|")]
            rows.append(cells)

        if len(rows) < 2:
            return

        header = rows[0]
        sep_idx = None
        for ri, row in enumerate(rows):
            if all(re.match(r'^[-:]+$', c) for c in row):
                sep_idx = ri
                break
        data = rows[sep_idx + 1:] if sep_idx is not None else rows[1:]
        ncols = len(header)
        col_w = min(effective_width / ncols, 45)
        col_widths = [col_w] * ncols
        remaining = effective_width - sum(col_widths)
        if remaining > 0 and ncols > 0:
            col_widths[-1] += remaining

        # Header
        pdf.set_font("Helvetica", "B", 8)
        pdf.set_fill_color(27, 42, 74)
        pdf.set_text_color(255, 255, 255)
        for ci, h in enumerate(header):
            pdf.cell(col_widths[ci], 6, h[:40], border=1, fill=True, new_x="END")
        pdf.ln()

        # Data
        for ri, row_data in enumerate(data):
            pdf.set_font("Helvetica", "", 8)
            pdf.set_text_color(51, 51, 51)
            if ri % 2 == 1:
                pdf.set_fill_color(242, 242, 242)
                fill = True
            else:
                pdf.set_fill_color(255, 255, 255)
                fill = True
            for ci in range(min(len(row_data), ncols)):
                val = re.sub(r'\*\*(.+?)\*\*', r'\1', row_data[ci])
                pdf.cell(col_widths[ci], 5.5, val[:50], border=1, fill=fill, new_x="END")
            pdf.ln()
        pdf.ln(3)

    while i < len(lines):
        line = lines[i].rstrip("\n")

        if skip_meta:
            if line.startswith("# ") or line.startswith("**") or line == "---" or line == "":
                i += 1
                continue
            if line.startswith("## "):
                skip_meta = False
            else:
                i += 1
                continue

        if in_table:
            if line.startswith("|"):
                table_buffer.append(line)
                i += 1
                continue
            else:
                render_table(table_buffer)
                table_buffer = []
                in_table = False

        if line.startswith("|") and "|" in line[1:]:
            in_table = True
            table_buffer = [line]
            i += 1
            continue

        if line == "---":
            i += 1
            continue

        if line.startswith("## "):
            heading = line[3:].strip()
            pdf.add_page()
            pdf.set_font("Helvetica", "B", 22)
            pdf.set_text_color(27, 42, 74)
            pdf.cell(0, 14, heading.upper(), new_x="LMARGIN", new_y="NEXT")
            pdf.set_draw_color(46, 117, 182)
            pdf.set_line_width(0.4)
            pdf.line(25, pdf.get_y(), pdf.w - 25, pdf.get_y())
            pdf.ln(8)
            i += 1
            continue

        if line.startswith("### Section "):
            heading = line[4:].strip()
            pdf.ln(6)
            pdf.set_font("Helvetica", "B", 13)
            pdf.set_text_color(46, 117, 182)
            pdf.cell(0, 8, heading, new_x="LMARGIN", new_y="NEXT")
            pdf.ln(3)
            i += 1
            continue

        if line.startswith("### "):
            heading = line[4:].strip()
            pdf.add_page()
            pdf.set_font("Helvetica", "B", 16)
            pdf.set_text_color(27, 42, 74)
            pdf.cell(0, 10, heading, new_x="LMARGIN", new_y="NEXT")
            pdf.set_draw_color(204, 204, 204)
            pdf.set_line_width(0.2)
            pdf.line(25, pdf.get_y(), pdf.w - 75, pdf.get_y())
            pdf.ln(5)
            i += 1
            continue

        if line.startswith("#### "):
            heading = line[5:].strip()
            pdf.ln(4)
            pdf.set_font("Helvetica", "B", 11)
            pdf.set_text_color(46, 117, 182)
            pdf.cell(0, 7, heading, new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)
            i += 1
            continue

        if line.startswith("- ") or line.startswith("* "):
            bullet_text = line[2:].strip()
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', bullet_text)
            clean = re.sub(r'\*(.+?)\*', r'\1', clean)
            clean = re.sub(r'---', '—', clean)
            write_paragraph(f"•  {clean}", indent=8)
            i += 1
            continue

        if line.strip() == "":
            i += 1
            continue

        write_paragraph(line)
        i += 1

    if in_table and table_buffer:
        render_table(table_buffer)

    # Appendices
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 22)
    pdf.set_text_color(27, 42, 74)
    pdf.cell(0, 14, "APPENDICES", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(8)

    pdf.set_font("Helvetica", "B", 14)
    pdf.set_text_color(46, 117, 182)
    pdf.cell(0, 8, "Appendix A — Release Notes Reference", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    ref_items = [
        ("Release Tag", "desk-bible-v1-publication"),
        ("Source Document", "release/desk-bible-v1/Desk_Bible_v1.md"),
        ("Release Notes", "release/desk-bible-v1/release_notes_v1.md"),
        ("Final Report", "release/desk-bible-v1/Desk_Bible_Final_Report.md"),
        ("Maintenance Guide", "release/desk-bible-v1/Maintenance_and_Governance_Guide.md"),
    ]
    for label, value in ref_items:
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_text_color(51, 51, 51)
        lw = pdf.get_string_width(f"{label}: ")
        pdf.cell(lw, 6, f"{label}: ", new_x="END")
        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(88, 88, 88)
        pdf.cell(0, 6, value, new_x="LMARGIN", new_y="NEXT")

    pdf.ln(6)
    pdf.set_font("Helvetica", "B", 14)
    pdf.set_text_color(46, 117, 182)
    pdf.cell(0, 8, "Appendix B — Quality Metrics Summary", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    metrics_tbl = [
        "| Metric | Value |",
        "|--------|-------|",
        "| Accuracy BLOCKERs | 0 |",
        "| Style MUST_FIX | 0 |",
        "| Broken Cross-References | 0 |",
        "| Publishing Failures | 0 |",
        "| Gate Pass Rate | 100% |",
        "| Retries (Batch 1+) | 0 |",
        "| Memory Reuse Ratio | 4.6:1 |",
        "| Tokens Per Product | ~14,866 |",
    ]
    render_table(metrics_tbl)

    pdf.output(out_path)
    print(f"PDF saved: {out_path}")
    print(f"PDF size: {os.path.getsize(out_path):,} bytes")
    print(f"PDF pages: {pdf.page_no()}")


if __name__ == "__main__":
    lines = parse_markdown(SRC)
    print(f"Source: {SRC}")
    print(f"Source lines: {len(lines)}")
    print()

    docx_path = os.path.join(OUT_DIR, "Desk_Bible_v1.docx")
    generate_docx(lines, docx_path)
    print()

    pdf_path = os.path.join(OUT_DIR, "Desk_Bible_v1.pdf")
    generate_pdf(lines, pdf_path)
